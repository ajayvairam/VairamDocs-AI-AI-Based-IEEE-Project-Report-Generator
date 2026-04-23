"""
DOCX export service — generates IEEE-formatted Word documents using python-docx.

This is a port of the client-side handleDocxDownload logic from App.tsx.
"""

import base64
import io
import re
from docx import Document
from docx.shared import Pt, Cm, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from app.models.schemas import ReportData, FormattingOptions


# ──────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────

def _parse_bold_markdown(paragraph, text: str, font_size: float = 10, is_italic: bool = False):
    """
    Parse simple **bold** markdown within text and add runs to a paragraph.
    Converts 'some **bold** text' into multiple runs with appropriate formatting.
    """
    if not text:
        return

    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(font_size)
            run.italic = is_italic
        else:
            run = paragraph.add_run(part)
            run.font.name = "Times New Roman"
            run.font.size = Pt(font_size)
            run.italic = is_italic


def _add_section_heading(doc: Document, text: str):
    """Add an IEEE-style section heading (uppercase, centered, bold)."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.space_before = Pt(12)
    para.space_after = Pt(6)
    run = para.add_run(text.upper())
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    return para


def _add_body_text(doc: Document, text: str, font_size: float = 10):
    """Add a justified body paragraph with bold-markdown parsing."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.space_after = Pt(6)
    _parse_bold_markdown(para, text or "", font_size=font_size)
    return para


def _set_columns(section, count: int = 2, spacing: float = 0.75):
    """Set a section to multi-column layout via XML manipulation."""
    if count < 2:
        return
    sectPr = section._sectPr
    cols = sectPr.find(qn('w:cols'))
    if cols is None:
        cols = parse_xml(f'<w:cols {nsdecls("w")} w:num="{count}" w:space="{int(spacing * 567)}"/>')
        sectPr.append(cols)
    else:
        cols.set(qn('w:num'), str(count))
        cols.set(qn('w:space'), str(int(spacing * 567)))


def _set_margins(section, top: float, bottom: float, left: float, right: float):
    """Set page margins in cm."""
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)


# ──────────────────────────────────────────────
# Main Export Function
# ──────────────────────────────────────────────

def generate_docx(report_data: ReportData, formatting: FormattingOptions) -> io.BytesIO:
    """
    Generate an IEEE-formatted DOCX document from report data.
    Returns a BytesIO buffer containing the DOCX file.
    """
    doc = Document()

    # ── Page Setup ──
    section = doc.sections[0]
    _set_margins(
        section,
        formatting.marginTop,
        formatting.marginBottom,
        formatting.marginLeft,
        formatting.marginRight,
    )
    if formatting.paperSize == "A4":
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
    else:  # Letter
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

    # ── HEADER SECTION (single column, centered) ──

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.space_after = Pt(12)
    title_run = title_para.add_run(report_data.title or "TITLE")
    title_run.bold = True
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(24)

    # Team Members
    for member in report_data.teamMembers:
        member_para = doc.add_paragraph()
        member_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        member_para.space_after = Pt(3)

        name_run = member_para.add_run(member.name or "Student Name")
        name_run.bold = True
        name_run.font.name = "Times New Roman"
        name_run.font.size = Pt(11)

        reg_run = member_para.add_run(f" ({member.registerNumber or 'Reg. No'})")
        reg_run.font.name = "Times New Roman"
        reg_run.font.size = Pt(10)

    # Department
    if report_data.department:
        dept_para = doc.add_paragraph()
        dept_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        dept_run = dept_para.add_run(report_data.department)
        dept_run.font.name = "Times New Roman"
        dept_run.font.size = Pt(10)
        dept_run.italic = True

    # College Name
    if report_data.collegeName:
        college_para = doc.add_paragraph()
        college_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        college_para.space_after = Pt(12)
        college_run = college_para.add_run(report_data.collegeName)
        college_run.font.name = "Times New Roman"
        college_run.font.size = Pt(10)

    # Guide
    if report_data.guideName:
        guide_para = doc.add_paragraph()
        guide_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        guide_para.space_after = Pt(20)

        guided_run = guide_para.add_run("Guided by: ")
        guided_run.italic = True
        guided_run.font.name = "Times New Roman"
        guided_run.font.size = Pt(11)

        name_run = guide_para.add_run(report_data.guideName)
        name_run.bold = True
        name_run.font.name = "Times New Roman"
        name_run.font.size = Pt(11)

    # ── Switch to double-column for body content ──
    # Add a section break and set columns on the new section
    new_section = doc.add_section()
    _set_margins(
        new_section,
        formatting.marginTop,
        formatting.marginBottom,
        formatting.marginLeft,
        formatting.marginRight,
    )
    if formatting.paperSize == "A4":
        new_section.page_width = Cm(21.0)
        new_section.page_height = Cm(29.7)
    else:
        new_section.page_width = Inches(8.5)
        new_section.page_height = Inches(11)

    if formatting.columns == 2:
        _set_columns(new_section, count=2, spacing=0.75)

    # Make this a continuous section break (no page break)
    new_section._sectPr.set(qn('w:type'), 'continuous')

    # ── BODY CONTENT ──

    # Abstract
    abstract_para = doc.add_paragraph()
    abstract_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    abstract_para.space_after = Pt(6)
    abs_label = abstract_para.add_run("Abstract—")
    abs_label.bold = True
    abs_label.italic = True
    abs_label.font.name = "Times New Roman"
    abs_label.font.size = Pt(9)
    _parse_bold_markdown(abstract_para, report_data.abstract, font_size=9, is_italic=True)

    # Keywords
    kw_para = doc.add_paragraph()
    kw_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    kw_para.space_after = Pt(12)
    kw_label = kw_para.add_run("Keywords—")
    kw_label.bold = True
    kw_label.italic = True
    kw_label.font.name = "Times New Roman"
    kw_label.font.size = Pt(9)
    _parse_bold_markdown(kw_para, report_data.keywords, font_size=9, is_italic=True)

    # Sections
    sections_content = [
        ("I. INTRODUCTION", report_data.introduction),
        ("II. PROBLEM STATEMENT", report_data.problemStatement),
        ("III. OBJECTIVES", report_data.objectives),
        ("IV. LITERATURE REVIEW", report_data.literatureReview),
        ("V. METHODOLOGY", report_data.methodology),
    ]

    for heading, content in sections_content:
        _add_section_heading(doc, heading)
        _add_body_text(doc, content)

    # Flowchart Image (after methodology)
    if report_data.flowchartImage:
        try:
            # Parse base64 data URI
            parts = report_data.flowchartImage.split(',')
            if len(parts) == 2:
                base64_data = parts[1]
            else:
                base64_data = parts[0]

            image_bytes = base64.b64decode(base64_data)
            image_stream = io.BytesIO(image_bytes)

            img_para = doc.add_paragraph()
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img_para.space_before = Pt(10)
            img_para.space_after = Pt(5)
            run = img_para.add_run()
            run.add_picture(image_stream, width=Inches(3.0))

            caption_para = doc.add_paragraph()
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_para.space_after = Pt(10)
            caption_run = caption_para.add_run("Fig. 1. Proposed System Architecture")
            caption_run.font.name = "Times New Roman"
            caption_run.font.size = Pt(8)
            caption_run.italic = True
            caption_run.bold = True

        except Exception as e:
            print(f"Warning: Could not add flowchart image: {e}")

    # Remaining sections
    remaining_sections = [
        ("VI. RESULTS", report_data.results),
        ("VII. CONCLUSION", report_data.conclusion),
        ("VIII. FUTURE SCOPE", report_data.futureScope),
    ]

    for heading, content in remaining_sections:
        _add_section_heading(doc, heading)
        _add_body_text(doc, content)

    # References
    _add_section_heading(doc, "REFERENCES")
    _add_body_text(doc, report_data.references, font_size=8)

    # ── Save to buffer ──
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
